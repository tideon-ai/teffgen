"""
Document ingestion pipeline.

Loads documents from many file formats, extracts metadata, deduplicates
content, and produces a list of `IngestedChunk` objects ready for indexing.

All third-party loaders (pymupdf, python-docx, bs4, ebooklib) are OPTIONAL.
Core formats (txt, md, json, jsonl, csv, html via stdlib) work with no
external dependencies.
"""

from __future__ import annotations

import csv
import hashlib
import json
import logging
import re
from collections.abc import Callable, Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from effgen.tools.builtin.retrieval import (
    Document,
    RecursiveChunker,
)

logger = logging.getLogger(__name__)


@dataclass
class IngestedChunk:
    """A chunk of an ingested document, ready for indexing."""

    id: str
    content: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)
    content_hash: str = ""

    def __post_init__(self):
        if not self.content_hash:
            self.content_hash = hashlib.sha256(
                self.content.encode("utf-8")
            ).hexdigest()[:16]

    def to_document(self) -> Document:
        md = dict(self.metadata)
        md.setdefault("source", self.source)
        md.setdefault("content_hash", self.content_hash)
        return Document(id=self.id, content=self.content, metadata=md)


# ---------------------------------------------------------------------------
# Loaders — always-available (stdlib only)
# ---------------------------------------------------------------------------

def _load_txt(path: Path) -> list[dict[str, Any]]:
    return [{
        "content": path.read_text(encoding="utf-8", errors="replace"),
        "metadata": {"type": "txt"},
    }]


def _load_markdown(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    # Attempt title extraction from first H1
    title = None
    for line in text.splitlines():
        m = re.match(r"^#\s+(.+)$", line.strip())
        if m:
            title = m.group(1).strip()
            break
    return [{
        "content": text,
        "metadata": {"type": "markdown", "title": title},
    }]


def _load_json(path: Path) -> list[dict[str, Any]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        out = []
        for i, item in enumerate(data):
            if isinstance(item, dict):
                content = item.get("content") or item.get("text") or json.dumps(item)
                meta = {k: v for k, v in item.items() if k not in ("content", "text")}
            else:
                content = str(item)
                meta = {}
            meta["type"] = "json"
            meta["index"] = i
            out.append({"content": content, "metadata": meta})
        return out
    if isinstance(data, dict):
        content = data.get("content") or data.get("text") or json.dumps(data)
        return [{"content": content, "metadata": {"type": "json"}}]
    return [{"content": str(data), "metadata": {"type": "json"}}]


def _load_jsonl(path: Path) -> list[dict[str, Any]]:
    out = []
    with open(path, encoding="utf-8") as f:
        for i, line in enumerate(f):
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(item, dict):
                content = item.get("content") or item.get("text") or json.dumps(item)
                meta = {k: v for k, v in item.items() if k not in ("content", "text")}
            else:
                content = str(item)
                meta = {}
            meta["type"] = "jsonl"
            meta["line"] = i
            out.append({"content": content, "metadata": meta})
    return out


def _load_csv(path: Path) -> list[dict[str, Any]]:
    out = []
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with open(path, encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            content = " ".join(str(v) for v in row.values() if v)
            if content.strip():
                out.append({
                    "content": content,
                    "metadata": {"type": "csv", "row": i, **row},
                })
    return out


def _load_html(path: Path) -> list[dict[str, Any]]:
    """Load HTML. Prefers bs4 if available, falls back to stdlib html.parser."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    title: str | None = None
    text: str

    try:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(raw, "html.parser")
        if soup.title and soup.title.string:
            title = soup.title.string.strip()
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text(separator="\n")
    except ImportError:
        from html.parser import HTMLParser

        class _Stripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts: list[str] = []
                self._skip = 0
                self._in_title = False
                self.title: str | None = None

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style"):
                    self._skip += 1
                if tag == "title":
                    self._in_title = True

            def handle_endtag(self, tag):
                if tag in ("script", "style") and self._skip:
                    self._skip -= 1
                if tag == "title":
                    self._in_title = False

            def handle_data(self, data):
                if self._skip:
                    return
                if self._in_title:
                    self.title = (self.title or "") + data
                self.parts.append(data)

        s = _Stripper()
        s.feed(raw)
        title = (s.title or "").strip() or None
        text = "\n".join(p for p in s.parts if p.strip())

    # Collapse whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return [{"content": text, "metadata": {"type": "html", "title": title}}]


# ---------------------------------------------------------------------------
# Optional loaders
# ---------------------------------------------------------------------------

def _load_pdf(path: Path) -> list[dict[str, Any]]:
    try:
        import fitz  # type: ignore  # pymupdf
    except ImportError as e:
        raise ImportError(
            "PDF support requires pymupdf. Install with: pip install pymupdf"
        ) from e

    doc = fitz.open(str(path))
    pages: list[str] = []
    metadata: dict[str, Any] = {"type": "pdf", "pages": doc.page_count}
    try:
        info = doc.metadata or {}
        for k in ("title", "author", "subject", "creationDate"):
            v = info.get(k)
            if v:
                metadata[k.lower().replace("creationdate", "date")] = v
    except Exception:
        pass
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return [{"content": "\n\n".join(pages), "metadata": metadata}]


def _load_docx(path: Path) -> list[dict[str, Any]]:
    try:
        import docx  # type: ignore  # python-docx
    except ImportError as e:
        raise ImportError(
            "DOCX support requires python-docx. Install with: pip install python-docx"
        ) from e

    d = docx.Document(str(path))
    paragraphs = [p.text for p in d.paragraphs if p.text]
    meta: dict[str, Any] = {"type": "docx"}
    try:
        core = d.core_properties
        if core.title:
            meta["title"] = core.title
        if core.author:
            meta["author"] = core.author
        if core.created:
            meta["date"] = str(core.created)
    except Exception:
        pass
    return [{"content": "\n\n".join(paragraphs), "metadata": meta}]


def _load_epub(path: Path) -> list[dict[str, Any]]:
    try:
        import ebooklib  # type: ignore
        from ebooklib import epub  # type: ignore
    except ImportError as e:
        raise ImportError(
            "EPUB support requires ebooklib. Install with: pip install ebooklib"
        ) from e

    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as e:
        raise ImportError(
            "EPUB support also requires beautifulsoup4. Install with: pip install beautifulsoup4"
        ) from e

    book = epub.read_epub(str(path))
    texts = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            texts.append(soup.get_text(separator="\n"))
    meta: dict[str, Any] = {"type": "epub"}
    try:
        titles = book.get_metadata("DC", "title")
        if titles:
            meta["title"] = titles[0][0]
        authors = book.get_metadata("DC", "creator")
        if authors:
            meta["author"] = authors[0][0]
    except Exception:
        pass
    return [{"content": "\n\n".join(texts), "metadata": meta}]


# Extension → loader map
LOADERS: dict[str, Callable[[Path], list[dict[str, Any]]]] = {
    ".txt": _load_txt,
    ".md": _load_markdown,
    ".markdown": _load_markdown,
    ".json": _load_json,
    ".jsonl": _load_jsonl,
    ".csv": _load_csv,
    ".tsv": _load_csv,
    ".html": _load_html,
    ".htm": _load_html,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".epub": _load_epub,
}


# ---------------------------------------------------------------------------
# Document Ingester
# ---------------------------------------------------------------------------

class DocumentIngester:
    """
    Ingest documents from a directory or list of paths, chunk them, and
    return `IngestedChunk` objects ready for indexing.

    Features:
    - Multi-format loaders (txt, md, json, jsonl, csv, html, pdf*, docx*, epub*)
    - Metadata extraction (title, author, date, source)
    - Content deduplication via SHA-256 hash of chunk text
    - Progress tracking via tqdm (optional)

    Optional loaders (*) raise a helpful ImportError at load time.
    """

    def __init__(
        self,
        chunker: Any = None,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        dedupe: bool = True,
        show_progress: bool = True,
    ):
        self.chunker = chunker or RecursiveChunker(
            chunk_size=chunk_size, overlap=chunk_overlap
        )
        self.dedupe = dedupe
        self.show_progress = show_progress
        self._seen_hashes: set[str] = set()

    def _iter_with_progress(self, items: list[Any], desc: str) -> Iterable[Any]:
        if not self.show_progress:
            return items
        try:
            from tqdm import tqdm  # type: ignore

            return tqdm(items, desc=desc)
        except ImportError:
            return items

    def supported_extensions(self) -> list[str]:
        return sorted(LOADERS.keys())

    def ingest(
        self,
        source: str | Path | list[str | Path],
        recursive: bool = True,
    ) -> list[IngestedChunk]:
        """
        Ingest documents from a file, directory, or list of paths.

        Args:
            source: file path, directory path, or list of paths.
            recursive: whether to recurse into subdirectories.

        Returns:
            A list of `IngestedChunk` objects.
        """
        paths: list[Path] = []
        if isinstance(source, (list, tuple)):
            for s in source:
                paths.extend(self._expand(Path(s), recursive))
        else:
            paths.extend(self._expand(Path(source), recursive))

        chunks: list[IngestedChunk] = []
        for path in self._iter_with_progress(paths, "Ingesting"):
            try:
                chunks.extend(self._ingest_file(path))
            except ImportError as e:
                logger.warning("Skipping %s: %s", path, e)
            except Exception as e:
                logger.warning("Failed to ingest %s: %s", path, e)

        logger.info("Ingested %d chunks from %d files", len(chunks), len(paths))
        return chunks

    def _expand(self, path: Path, recursive: bool) -> list[Path]:
        if path.is_file():
            return [path]
        if path.is_dir():
            pattern = "**/*" if recursive else "*"
            return [
                p for p in path.glob(pattern)
                if p.is_file() and p.suffix.lower() in LOADERS
            ]
        logger.warning("Path does not exist: %s", path)
        return []

    def _ingest_file(self, path: Path) -> list[IngestedChunk]:
        ext = path.suffix.lower()
        loader = LOADERS.get(ext)
        if loader is None:
            return []

        docs = loader(path)
        if not docs:
            return []

        out: list[IngestedChunk] = []
        source = str(path)
        base_id = path.stem

        for doc_idx, doc in enumerate(docs):
            content = doc.get("content", "") or ""
            if not content.strip():
                continue
            doc_meta = doc.get("metadata", {}) or {}

            # Chunk the content
            doc_id = f"{base_id}_{doc_idx}" if len(docs) > 1 else base_id
            chunk_docs = self.chunker.chunk(content, doc_id)

            for c in chunk_docs:
                if not c.content.strip():
                    continue
                content_hash = hashlib.sha256(
                    c.content.encode("utf-8")
                ).hexdigest()[:16]

                if self.dedupe and content_hash in self._seen_hashes:
                    continue
                self._seen_hashes.add(content_hash)

                merged_meta = {**doc_meta, **c.metadata}
                merged_meta["source"] = source
                out.append(IngestedChunk(
                    id=c.id,
                    content=c.content,
                    source=source,
                    metadata=merged_meta,
                    content_hash=content_hash,
                ))

        return out

    def reset_dedupe(self) -> None:
        """Clear the deduplication cache."""
        self._seen_hashes.clear()
